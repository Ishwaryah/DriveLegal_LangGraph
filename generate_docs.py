import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_submission_doc():
    doc = Document()
    
    # Title formatting
    title = doc.add_heading('DriveLegal', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0, 85, 113) # A nice dark blue/teal
    
    subtitle = doc.add_paragraph('Project Submission Document')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('1. Project Overview', level=1)
    doc.add_paragraph(
        "DriveLegal is a cutting-edge, AI-powered mobile application designed to simplify "
        "traffic laws, regulations, and fine management. Built with a focus on ease of use "
        "and professional aesthetics, it serves as a digital legal companion for every driver. "
        "The project aims to make roads safer and traffic laws more accessible to the general public."
    )
    
    doc.add_heading('2. Key Features', level=1)
    features = [
        "AI Legal Chatbot: Instant, accurate answers to complex traffic law queries powered by advanced NLP.",
        "Real-time Geofencing: Location-aware traffic zone alerts and zone-specific regulations.",
        "Challan Calculator: Quickly check vehicle registration numbers to calculate pending fines and view violation history.",
        "Voice Integration: Hands-free interactions with voice-to-text search for increased safety while driving.",
        "Premium Interface: A government-branded, high-performance UI designed for clarity and professional appeal.",
        "Offline-First: Local SQLite database integration for fast access to traffic rules even without internet connectivity."
    ]
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')
        
    doc.add_heading('3. Technology Stack', level=1)
    
    doc.add_heading('Frontend (Mobile)', level=2)
    mobile_stack = [
        "Framework: Expo / React Native",
        "Language: TypeScript",
        "Navigation: Expo Router (v3)",
        "Data Fetching: TanStack Query (React Query)",
        "Maps: MapLibre",
        "Storage: Expo SQLite & Async Storage"
    ]
    for tech in mobile_stack:
        doc.add_paragraph(tech, style='List Bullet 2')
        
    doc.add_heading('Backend', level=2)
    backend_stack = [
        "Framework: FastAPI (Python)",
        "Server: Uvicorn",
        "Validation: Pydantic",
        "NLP: Custom NLP Pipeline"
    ]
    for tech in backend_stack:
        doc.add_paragraph(tech, style='List Bullet 2')
        
    doc.add_heading('Data & Scraper', level=2)
    data_stack = [
        "Scraper: Python-based automation for up-to-date fine schedules.",
        "Database: Ported SQLite for mobile consumption."
    ]
    for tech in data_stack:
        doc.add_paragraph(tech, style='List Bullet 2')

    doc.add_heading('4. Project Structure', level=1)
    doc.add_paragraph(
        "DriveLegal/\n"
        "├── mobile/             # React Native / Expo Application\n"
        "│   ├── app/            # Expo Router main app logic\n"
        "│   ├── components/     # Reusable UI components\n"
        "│   ├── hooks/          # Custom React hooks\n"
        "│   └── local_db/       # SQLite database management\n"
        "├── backend/            # FastAPI Server\n"
        "│   ├── modules/        # NLP, Fines, and Geofencing logic\n"
        "│   └── data/           # Reference rules and zone data\n"
        "└── scraper/            # Data collection scripts"
    )
    
    doc.add_page_break()
    doc.add_heading('5. System Installation & Setup', level=1)
    doc.add_paragraph("Prerequisites: Node.js (v18+), Python (3.9+), Expo Go app")
    
    doc.add_heading('Backend Setup:', level=2)
    doc.add_paragraph("1. Navigate to the backend directory: cd backend\n"
                      "2. Install dependencies: pip install -r requirements.txt\n"
                      "3. Start the server: python main.py")
                      
    doc.add_heading('Mobile App Setup:', level=2)
    doc.add_paragraph("1. Navigate to the mobile directory: cd mobile\n"
                      "2. Install dependencies: npm install\n"
                      "3. Start the Expo development server: npx expo start")
                      
    doc.add_heading('6. Disclaimer', level=1)
    doc.add_paragraph(
        "DriveLegal is an educational tool and does not constitute official legal advice. "
        "Always consult with legal professionals or relevant government authorities for official "
        "traffic law interpretations."
    )
    
    doc.save('DriveLegal_Project_Submission.docx')
    print("Submission document generated successfully.")

if __name__ == '__main__':
    create_submission_doc()
